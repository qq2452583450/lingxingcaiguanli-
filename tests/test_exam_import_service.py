from pathlib import Path
import importlib

import pytest
from docx import Document

from services.exam_import_service import (
    BUNDLED_FORMAL_EXAM_POOL_DIR,
    DESKTOP_QUESTION_BANK_DIR,
    FORMAL_EXAM_POOL_SOURCE_TYPE,
    ensure_exam_sources_imported,
    get_question_bank_dir,
    import_exam_papers_from_question_bank_dir,
    insert_paper,
    parse_exam_docx,
    parse_formal_exam_pool_docx,
    parse_literature_question_bank_docx,
    replace_exam_paper_from_question_bank_docx,
    sync_missing_question_bank_papers,
    sync_formal_exam_pool_papers,
    sync_question_bank_reference_answers,
    import_exam_papers_from_docx,
)
from services.exam_service import (
    get_current_exam_paper,
    get_paper_questions,
    get_random_practice_questions,
    list_papers,
)


def source_docx():
    return next(Path("docs/exam_sources").glob("*\u6750\u6599\u8fdb\u573a\u9a8c\u6536\u6807\u51c6\u4e13\u9879\u8003\u8bd5\u5377*.docx"))


def test_parse_receiving_exam_docx_has_five_complete_papers():
    papers = parse_exam_docx(source_docx())

    assert [paper["title"] for paper in papers] == [
        "\u7b2c\u4e00\u5957\uff08\u65b0\u7f16\u5b9e\u64cd\u7248\uff09",
        "\u7b2c\u4e8c\u5957\uff08\u65b0\u7f16\u6848\u4f8b\u7248\uff09",
        "\u7b2c\u4e09\u5957\uff08\u65b0\u7f16\u5185\u63a7\u7248\uff09",
        "\u7b2c\u56db\u5957\uff08\u65b0\u7f16\u5b9e\u64cd\u6613\u9519\u7248\uff09",
        "\u7b2c\u4e94\u5957\uff08\u65b0\u7f16\u7efc\u5408\u62bc\u9898\u7248\uff09",
    ]
    assert all(len(paper["questions"]) == 38 for paper in papers)


def test_parse_formal_exam_pool_has_three_unique_60_minute_100_point_papers():
    papers_by_title = {}
    for source in sorted(BUNDLED_FORMAL_EXAM_POOL_DIR.glob("*.docx")):
        paper = parse_formal_exam_pool_docx(source)
        papers_by_title.setdefault(paper["title"], paper)

    assert set(papers_by_title) == {
        "综合题库一（满分 100 分）",
        "综合题库二（满分 100 分）",
        "综合题库三（满分 100 分）",
    }
    all_stems = []
    for paper in papers_by_title.values():
        assert paper["duration_minutes"] == 60
        assert paper["total_score"] == 100
        assert len(paper["questions"]) == 35
        all_stems.extend(question["stem"] for question in paper["questions"])
        assert sum(question["score"] for question in paper["questions"]) == 100
        assert {
            question_type: sum(
                question["question_type"] == question_type
                for question in paper["questions"]
            )
            for question_type in ("single_choice", "multiple_choice", "true_false")
        } == {"single_choice": 20, "multiple_choice": 10, "true_false": 5}
    assert len(all_stems) == len(set(all_stems))


def test_sync_formal_exam_pool_is_idempotent_and_excludes_it_from_practice(test_db):
    first = sync_formal_exam_pool_papers()
    second = sync_formal_exam_pool_papers()
    papers = list_papers()

    assert first == {
        "inserted": 3, "archived": 0, "retired": 0, "unchanged": 0, "source_files": 3
    }
    assert second == {
        "inserted": 0, "archived": 0, "retired": 0, "unchanged": 3, "source_files": 3
    }
    assert [paper["title"] for paper in papers] == [
        "综合题库一（满分 100 分）",
        "综合题库二（满分 100 分）",
        "综合题库三（满分 100 分）",
    ]
    assert {paper["source_type"] for paper in papers} == {FORMAL_EXAM_POOL_SOURCE_TYPE}
    assert get_random_practice_questions(limit=100) == []
    assert get_random_practice_questions(limit=100, paper_id=papers[0]["id"]) == []


def test_imported_choice_questions_keep_source_option_e(test_db):
    paper_id = insert_paper(
        test_db.cursor(),
        {
            "title": "With E option paper",
            "duration_minutes": 60,
            "questions": [
                {
                    "question_type": "multiple_choice",
                    "order_no": 1,
                    "stem": "Multiple choice",
                    "correct_answer": "ABCE",
                    "reference_answer": "Explanation",
                    "keywords": "",
                    "score": 3,
                    "options": [
                        {"key": "A", "text": "Option A"},
                        {"key": "B", "text": "Option B"},
                        {"key": "C", "text": "Option C"},
                        {"key": "D", "text": "Option D"},
                        {"key": "E", "text": "Option E"},
                    ],
                }
            ],
        },
    )
    test_db.commit()

    question = get_paper_questions(paper_id)[0]

    assert question["correct_answer"] == "ABCE"
    assert [option["key"] for option in question["options"]] == ["A", "B", "C", "D", "E"]


def test_imported_choice_questions_drop_answers_without_matching_options(test_db):
    paper_id = insert_paper(
        test_db.cursor(),
        {
            "title": "No E option paper",
            "duration_minutes": 60,
            "questions": [
                {
                    "question_type": "multiple_choice",
                    "order_no": 1,
                    "stem": "Multiple choice",
                    "correct_answer": "ABCE",
                    "reference_answer": "Explanation",
                    "keywords": "",
                    "score": 3,
                    "options": [
                        {"key": "A", "text": "Option A"},
                        {"key": "B", "text": "Option B"},
                        {"key": "C", "text": "Option C"},
                        {"key": "D", "text": "Option D"},
                    ],
                }
            ],
        },
    )
    test_db.commit()

    question = get_paper_questions(paper_id)[0]

    assert question["correct_answer"] == "ABC"
    assert [option["key"] for option in question["options"]] == ["A", "B", "C", "D"]


def test_literature_question_bank_uses_section_scores(tmp_path):
    source_path = tmp_path / "scored-bank.docx"
    doc = Document()
    doc.add_paragraph("综合题库四（满分100分）")
    doc.add_paragraph("一、单项选择题（每题 3 分）")
    doc.add_paragraph("1. 单选题？")
    doc.add_paragraph("A. 甲 B. 乙 C. 丙 D. 丁")
    doc.add_paragraph("答案：A")
    doc.add_paragraph("二、多项选择题（每题 3 分）")
    doc.add_paragraph("1. 多选题？")
    doc.add_paragraph("A. 甲 B. 乙 C. 丙 D. 丁 E. 戊")
    doc.add_paragraph("答案：AB")
    doc.add_paragraph("三、判断题（每题 2 分）")
    doc.add_paragraph("1. 判断题。")
    doc.add_paragraph("答案：√")
    doc.save(source_path)

    paper = parse_literature_question_bank_docx(source_path)

    assert [question["score"] for question in paper["questions"]] == [3, 3, 2]
    assert paper["total_score"] == 8


def test_sync_formal_exam_pool_preserves_current_exam_and_daily_practice_bank(test_db):
    import_exam_papers_from_docx(source_docx())
    current = list_papers()[0]
    test_db.execute(
        "INSERT INTO exam_settings (key, value) VALUES (?, ?)",
        ("current_exam_paper_id", str(current["id"])),
    )

    sync_formal_exam_pool_papers()
    practice = get_random_practice_questions(limit=200)

    assert get_current_exam_paper()["id"] == current["id"]
    assert practice
    assert {
        "综合题库一（满分 100 分）",
        "综合题库二（满分 100 分）",
        "综合题库三（满分 100 分）",
    }.isdisjoint({question["paper_title"] for question in practice})


def test_sync_formal_exam_pool_archives_retired_papers_and_clears_old_current(test_db):
    cursor = test_db.cursor()
    retired_ids = []
    for title in (
        "项目物资管理综合考核试卷（第一套）",
        "项目物资管理综合考核试卷（第二套）",
    ):
        cursor.execute(
            """
            INSERT INTO exam_papers (
                title, duration_minutes, total_score, source_type, create_time
            ) VALUES (?, ?, ?, ?, ?)
            """,
            (title, 60, 100, "exam", "2026-07-29 00:00:00"),
        )
        retired_ids.append(cursor.lastrowid)
    cursor.execute(
        "INSERT INTO exam_settings (key, value) VALUES (?, ?)",
        ("current_exam_paper_id", str(retired_ids[0])),
    )
    test_db.commit()

    result = sync_formal_exam_pool_papers()

    retired = test_db.execute(
        "SELECT id, source_type FROM exam_papers WHERE id IN (?, ?) ORDER BY id",
        retired_ids,
    ).fetchall()
    assert result["retired"] == 2
    assert [row["source_type"] for row in retired] == ["archived_exam", "archived_exam"]
    assert get_current_exam_paper() is None
    assert len(list_papers()) == 3


def test_import_exam_papers_does_not_auto_select_current_exam_paper(test_db):
    result = import_exam_papers_from_docx(source_docx())

    papers = list_papers()
    current = get_current_exam_paper()

    assert result == {"inserted": 5, "removed": 0}
    assert len(papers) == 5
    assert current is None
    assert len(get_paper_questions(papers[0]["id"])) == 38


def test_desktop_question_bank_dir_points_to_current_user_folder():
    assert DESKTOP_QUESTION_BANK_DIR.name == "\u9898\u5e93"
    assert DESKTOP_QUESTION_BANK_DIR.parent.name == "Desktop"
    assert get_question_bank_dir(DESKTOP_QUESTION_BANK_DIR) == DESKTOP_QUESTION_BANK_DIR


def test_parse_desktop_question_bank_docx_keeps_inline_explanations(tmp_path):
    source = Path.home() / "Desktop" / "\u9898\u5e93" / "\u4e2d\u5929\u5efa\u8bbe\u6210\u672c\u5e73\u53f0\u7269\u8d44\u7ba1\u7406V4.0\u64cd\u4f5c\u9898\u5e93.docx"
    if not source.exists():
        pytest.skip("Desktop question bank is not available on this machine")

    paper = parse_literature_question_bank_docx(source)

    assert paper["questions"]
    assert all(question["reference_answer"] for question in paper["questions"])
    assert "\u624b\u518c\u7b2c\u4e8c\u7ae0" in paper["questions"][0]["reference_answer"]


def test_sync_question_bank_reference_answers_backfills_existing_questions(test_db, tmp_path):
    source_dir = Path.home() / "Desktop" / "\u9898\u5e93"
    if not source_dir.exists():
        pytest.skip("Desktop question bank is not available on this machine")

    import_exam_papers_from_question_bank_dir(source_dir)
    question = test_db.execute(
        """
        SELECT q.id, q.reference_answer
        FROM exam_questions q
        JOIN exam_papers p ON p.id = q.paper_id
        WHERE p.title = ?
        ORDER BY q.id
        LIMIT 1
        """,
        ("\u4e2d\u5929\u5efa\u8bbe\u6210\u672c\u5e73\u53f0\u7269\u8d44\u7ba1\u7406V4.0\u64cd\u4f5c\u9898\u5e93",),
    ).fetchone()
    assert question["reference_answer"]
    test_db.execute("UPDATE exam_questions SET reference_answer = '' WHERE id = ?", (question["id"],))

    result = sync_question_bank_reference_answers(source_dir)
    updated = test_db.execute(
        "SELECT reference_answer FROM exam_questions WHERE id = ?",
        (question["id"],),
    ).fetchone()

    assert result["updated"] >= 1
    assert "\u624b\u518c\u7b2c\u4e8c\u7ae0" in updated["reference_answer"]


def test_sync_question_bank_reference_answers_can_match_by_options_when_stem_changed(test_db):
    source_dir = Path.home() / "Desktop" / "\u9898\u5e93"
    if not source_dir.exists():
        pytest.skip("Desktop question bank is not available on this machine")

    import_exam_papers_from_question_bank_dir(source_dir)
    question = test_db.execute(
        """
        SELECT q.id, q.reference_answer
        FROM exam_questions q
        JOIN exam_papers p ON p.id = q.paper_id
        WHERE p.title = ?
        ORDER BY q.id
        LIMIT 1
        """,
        ("\u4e2d\u5929\u5efa\u8bbe\u6210\u672c\u5e73\u53f0\u7269\u8d44\u7ba1\u7406V4.0\u64cd\u4f5c\u9898\u5e93",),
    ).fetchone()
    assert question["reference_answer"]
    test_db.execute(
        "UPDATE exam_questions SET stem = ?, reference_answer = '' WHERE id = ?",
        ("\u6539\u5199\u540e\u7684\u9898\u5e72\uff0c\u4f46\u9009\u9879\u4e00\u81f4", question["id"]),
    )

    result = sync_question_bank_reference_answers(source_dir)
    updated = test_db.execute(
        "SELECT reference_answer FROM exam_questions WHERE id = ?",
        (question["id"],),
    ).fetchone()

    assert result["updated"] >= 1
    assert "\u624b\u518c\u7b2c\u4e8c\u7ae0" in updated["reference_answer"]


def test_ensure_exam_sources_imported_is_idempotent(test_db):
    first = ensure_exam_sources_imported()
    second = ensure_exam_sources_imported()

    assert first == {"created": True, "paper_count": 9}
    assert second == {"created": False, "paper_count": 9}
    assert len([paper for paper in list_papers() if paper["source_type"] == "exam"]) == 6
    assert len(
        [paper for paper in list_papers() if paper["source_type"] == FORMAL_EXAM_POOL_SOURCE_TYPE]
    ) == 3


def test_ensure_exam_sources_imported_keeps_missing_current_setting(test_db):
    import_exam_papers_from_docx(source_docx())
    test_db.execute(
        "DELETE FROM exam_settings WHERE key = ?",
        ("current_exam_paper_id",),
    )

    result = ensure_exam_sources_imported()

    assert result == {"created": True, "paper_count": 9}
    assert get_current_exam_paper() is None


def test_ensure_exam_sources_imported_does_not_replace_non_exam_current_setting(test_db):
    import_exam_papers_from_docx(source_docx())
    cursor = test_db.cursor()
    cursor.execute(
        """
        INSERT INTO exam_papers (
            title, duration_minutes, total_score, source_type, create_time
        ) VALUES (?, ?, ?, ?, ?)
        """,
        ("\u9898\u5e93\u53c2\u8003\u5377", 50, 100, "bank", "2026-06-27 00:00:00"),
    )
    bank_paper_id = cursor.lastrowid
    cursor.execute(
        "INSERT OR REPLACE INTO exam_settings (key, value) VALUES (?, ?)",
        ("current_exam_paper_id", str(bank_paper_id)),
    )
    test_db.commit()

    result = ensure_exam_sources_imported()

    assert result == {"created": True, "paper_count": 9}
    assert get_current_exam_paper() is None


def test_ensure_exam_sources_imported_does_not_replace_stale_current_setting(test_db):
    import_exam_papers_from_docx(source_docx())
    test_db.execute(
        "INSERT OR REPLACE INTO exam_settings (key, value) VALUES (?, ?)",
        ("current_exam_paper_id", "999999"),
    )

    result = ensure_exam_sources_imported()

    assert result == {"created": True, "paper_count": 9}
    assert get_current_exam_paper() is None
    assert len([paper for paper in list_papers() if paper["source_type"] == "exam"]) == 6


def test_sync_missing_question_bank_papers_adds_only_absent_titles(test_db):
    source_dir = Path("docs/exam_sources/question_bank")
    initial_files = sorted(source_dir.glob("*.docx"))
    first_path = initial_files[0]
    first_paper = parse_literature_question_bank_docx(first_path)
    insert_paper(test_db.cursor(), first_paper, source_type="exam")
    test_db.commit()

    result = sync_missing_question_bank_papers(source_dir)
    second = sync_missing_question_bank_papers(source_dir)
    papers = list_papers()

    assert result == {
        "inserted": len(initial_files) - 1,
        "archived": 0,
        "migrated_attempts": 0,
        "skipped": 0,
        "source_files": len(initial_files),
    }
    assert second == {
        "inserted": 0,
        "archived": 0,
        "migrated_attempts": 0,
        "skipped": 0,
        "source_files": len(initial_files),
    }
    assert len([paper for paper in papers if paper["source_type"] == "exam"]) == len(initial_files)
    assert len([paper for paper in papers if paper["title"] == first_paper["title"]]) == 1


def test_sync_missing_question_bank_papers_refreshes_changed_title(test_db):
    source_dir = Path("docs/exam_sources/question_bank")
    source_path = source_dir / "综合题库四（满分100分）.docx"
    if not source_path.exists():
        pytest.skip("Fourth comprehensive bank source is not available")
    paper = parse_literature_question_bank_docx(source_path)
    paper_without_e = {
        **paper,
        "questions": [
            {
                **question,
                "options": [
                    option for option in question.get("options", []) if option.get("key") != "E"
                ],
                "correct_answer": str(question.get("correct_answer", "")).replace("E", ""),
            }
            for question in paper["questions"]
        ],
    }
    insert_paper(test_db.cursor(), paper_without_e, source_type="exam")
    test_db.commit()

    result = sync_missing_question_bank_papers(source_dir)
    active_paper = next(p for p in list_papers() if p["title"] == paper["title"])
    questions = get_paper_questions(active_paper["id"])

    assert result["inserted"] == len(list(source_dir.glob("*.docx")))
    assert result["archived"] == 1
    assert result["migrated_attempts"] == 0
    assert any(option["key"] == "E" for question in questions for option in question["options"])
    assert any("E" in question["correct_answer"] for question in questions)


def test_sync_missing_question_bank_papers_moves_active_attempts_to_refreshed_paper(test_db):
    source_dir = Path("docs/exam_sources/question_bank")
    source_path = source_dir / "综合题库四（满分100分）.docx"
    if not source_path.exists():
        pytest.skip("Fourth comprehensive bank source is not available")
    paper = parse_literature_question_bank_docx(source_path)
    paper_without_e = {
        **paper,
        "questions": [
            {
                **question,
                "options": [
                    option for option in question.get("options", []) if option.get("key") != "E"
                ],
                "correct_answer": str(question.get("correct_answer", "")).replace("E", ""),
            }
            for question in paper["questions"]
        ],
    }
    old_paper_id = insert_paper(test_db.cursor(), paper_without_e, source_type="exam")
    user_id = 1
    test_db.execute(
        """
        INSERT INTO users (id, username, password, real_name, create_time)
        VALUES (?, ?, ?, ?, ?)
        """,
        (user_id, "active_old_paper", "password", "旧卷开考人员", "2026-09-01 00:00:00"),
    )
    test_db.execute(
        """
        INSERT INTO exam_attempts (user_id, paper_id, status, started_at)
        VALUES (?, ?, 'in_progress', ?)
        """,
        (user_id, old_paper_id, "2026-09-01 00:00:00"),
    )
    test_db.commit()

    result = sync_missing_question_bank_papers(source_dir)
    attempt = test_db.execute(
        "SELECT paper_id FROM exam_attempts WHERE user_id = ? AND status = 'in_progress'",
        (user_id,),
    ).fetchone()
    active_paper = next(p for p in list_papers() if p["title"] == paper["title"])
    questions = get_paper_questions(active_paper["id"])

    assert result["migrated_attempts"] == 1
    assert attempt["paper_id"] == active_paper["id"]
    assert attempt["paper_id"] != old_paper_id
    assert any(option["key"] == "E" for question in questions for option in question["options"])


def test_startup_exam_source_hook_uses_import_helper(monkeypatch):
    calls = []

    def fake_ensure_exam_sources_imported():
        calls.append(True)
        return {"created": False, "paper_count": 5}

    monkeypatch.setenv("SECRET_KEY", "test-secret-key")
    app_module = importlib.import_module("app")
    monkeypatch.setattr(
        "services.exam_import_service.ensure_exam_sources_imported",
        fake_ensure_exam_sources_imported,
    )

    assert app_module.ensure_startup_exam_sources() == {"created": False, "paper_count": 5}
    assert calls == [True]


def test_reimport_archives_old_formal_exam_without_removing_lifecycle_rows(test_db):
    first_result = import_exam_papers_from_docx(source_docx())
    old_current = list_papers()[0]
    old_question = get_paper_questions(old_current["id"])[0]

    cursor = test_db.cursor()
    cursor.execute(
        """
        INSERT INTO users (username, password, real_name, create_time)
        VALUES (?, ?, ?, ?)
        """,
        ("exam_user", "password", "Exam User", "2026-06-26 00:00:00"),
    )
    user_id = cursor.lastrowid
    cursor.execute(
        """
        INSERT INTO exam_attempts (
            user_id, paper_id, status, started_at
        ) VALUES (?, ?, ?, ?)
        """,
        (user_id, old_current["id"], "completed", "2026-06-26 00:00:00"),
    )
    attempt_id = cursor.lastrowid
    cursor.execute(
        """
        INSERT INTO exam_answers (
            attempt_id, question_id, answer_text, auto_score
        ) VALUES (?, ?, ?, ?)
        """,
        (attempt_id, old_question["id"], "A", 0),
    )
    answer_id = cursor.lastrowid
    cursor.execute(
        """
        INSERT INTO exam_subjective_reviews (
            answer_id, reviewer_id, suggested_score, final_score, reviewed_at
        ) VALUES (?, ?, ?, ?, ?)
        """,
        (answer_id, user_id, 0, 0, "2026-06-26 00:00:00"),
    )
    review_id = cursor.lastrowid
    cursor.execute(
        """
        INSERT INTO exam_practice_attempts (
            user_id, question_id, answer_text, is_correct, created_at
        ) VALUES (?, ?, ?, ?, ?)
        """,
        (user_id, old_question["id"], "A", 1, "2026-06-26 00:00:00"),
    )
    practice_id = cursor.lastrowid
    test_db.commit()

    second_result = import_exam_papers_from_docx(source_docx())
    papers = list_papers()
    current = get_current_exam_paper()

    assert first_result == {"inserted": 5, "removed": 0}
    assert second_result == {"inserted": 5, "archived": 5, "removed": 0}
    assert len([paper for paper in papers if paper["source_type"] == "exam"]) == 5
    assert current is None
    assert test_db.execute(
        "SELECT COUNT(*) FROM exam_attempts WHERE id = ?",
        (attempt_id,),
    ).fetchone()[0] == 1
    assert test_db.execute(
        "SELECT COUNT(*) FROM exam_answers WHERE id = ?",
        (answer_id,),
    ).fetchone()[0] == 1
    assert test_db.execute(
        "SELECT COUNT(*) FROM exam_subjective_reviews WHERE id = ?",
        (review_id,),
    ).fetchone()[0] == 1
    assert test_db.execute(
        "SELECT COUNT(*) FROM exam_practice_attempts WHERE id = ?",
        (practice_id,),
    ).fetchone()[0] == 1
    assert test_db.execute(
        "SELECT source_type FROM exam_papers WHERE id = ?",
        (old_current["id"],),
    ).fetchone()[0] == "archived_exam"


def test_replace_question_bank_archives_only_matching_active_paper(test_db):
    source_dir = Path("docs/exam_sources/question_bank")
    replacement_path = source_dir / "物资管理-结算单调差应用规范题库.docx"
    import_exam_papers_from_question_bank_dir(source_dir)
    before = list_papers()
    before_by_title = {paper["title"]: paper for paper in before}
    target_title = "物资管理-结算单调差应用规范题库"
    target_id = before_by_title[target_title]["id"]

    result = replace_exam_paper_from_question_bank_docx(replacement_path)

    after = list_papers()
    after_by_title = {paper["title"]: paper for paper in after}
    assert {key: result[key] for key in ("inserted", "archived", "removed")} == {
        "inserted": 1,
        "archived": 1,
        "removed": 0,
    }
    assert set(after_by_title) == set(before_by_title)
    assert after_by_title[target_title]["id"] != target_id
    assert len(get_paper_questions(after_by_title[target_title]["id"])) == 30
    assert test_db.execute(
        "SELECT source_type FROM exam_papers WHERE id = ?", (target_id,)
    ).fetchone()[0] == "archived_exam"
